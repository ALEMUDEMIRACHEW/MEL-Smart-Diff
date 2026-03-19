import streamlit as st
from google import genai
from google.genai import types
import fitz  # PyMuPDF
from docx import Document
import io
import os
import shutil
import pandas as pd
from datetime import datetime
from thefuzz import fuzz
from redlines import Redlines

# --- 1. CONFIGURATION & STYLING ---
st.set_page_config(page_title="Smart-Diff Pro", layout="wide", page_icon="🔍")

# Consistent Styling for 2026 UI
st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    ins { background-color: #d4edda; text-decoration: none; color: #155724; padding: 2px; border-radius: 3px; }
    del { background-color: #f8d7da; color: #721c24; padding: 2px; border-radius: 3px; }
    .stMetric { background-color: white; padding: 15px; border-radius: 10px; border: 1px solid #eee; box-shadow: 0 2px 4px rgba(0,0,0,0.05); }
    .keyword-alert { background-color: #fff3cd; color: #856404; padding: 12px; border-radius: 5px; border: 1px solid #ffeeba; margin-bottom: 10px; font-weight: bold; }
    </style>
    """, unsafe_allow_html=True)

# API & Security
api_key = st.secrets.get("GEMINI_API_KEY", "YOUR_KEY_HERE")
APP_PASSWORD = st.secrets.get("APP_PASSWORD", "admin123")

# Persistent Session State
if "history" not in st.session_state: st.session_state.history = []
if "batch_log" not in st.session_state: st.session_state.batch_log = []

# --- 2. CORE UTILITIES (Verified) ---
def open_local_file(path):
    """Feature: Windows Local Integration."""
    try:
        if os.path.exists(path):
            os.startfile(path)
        else: st.error(f"Path no longer valid: {path}")
    except Exception as e: st.error(f"System error: {e}")

def archive_file(source_path):
    """Feature: Automatic Workflow Archiving."""
    try:
        folder = os.path.dirname(source_path)
        archive_dir = os.path.join(folder, "Audited_Results")
        if not os.path.exists(archive_dir):
            os.makedirs(archive_dir)
        dest_path = os.path.join(archive_dir, os.path.basename(source_path))
        shutil.move(source_path, dest_path)
        return True
    except Exception as e:
        st.error(f"Archive failed: {e}")
        return False

def extract_text(source):
    """Feature: Multi-Format Extraction (PDF/DOCX)."""
    try:
        if isinstance(source, str):
            if source.lower().endswith('.pdf'):
                doc = fitz.open(source)
                return "\n".join([p.get_text() for p in doc])
            elif source.lower().endswith('.docx'):
                doc = Document(source)
                return "\n".join([p.text for p in doc.paragraphs])
        else:
            content = source.getvalue()
            if source.name.lower().endswith('.pdf'):
                doc = fitz.open(stream=content, filetype="pdf")
                return "\n".join([p.get_text() for p in doc])
            elif source.name.lower().endswith('.docx'):
                doc = Document(io.BytesIO(content))
                return "\n".join([p.text for p in doc.paragraphs])
        return ""
    except: return "Extraction Error"

def run_audit(text_a, text_b, key, mode):
    """Feature: AI Core (Gemini 3 Flash)."""
    client = genai.Client(api_key=key)
    config = types.GenerateContentConfig(
        system_instruction=f"Aviation Auditor Mode: {mode}. Compare original and revised maintenance task cards. Flag 🟢[ADD], 🔴[DEL], 🟠[MOD]. Use 5-word anchors.",
        temperature=0.0
    )
    prompt = f"ORIGINAL:\n{text_a}\n\nREVISED:\n{text_b}"
    return client.models.generate_content(model="gemini-3-flash", contents=prompt, config=config).text

# --- 3. SIDEBAR (Full Feature Restore) ---
with st.sidebar:
    st.header("🔐 Access Control")
    pwd_input = st.text_input("App Password", type="password")
    st.divider()
    st.header("⚙️ Audit Profile")
    focus_mode = st.radio("Sensitivity:", ["Strict Audit", "Logic Only", "Summary"])
    st.header("🚨 Safety Keywords")
    keywords = st.text_input("Watchlist:", "Caution, Warning, Note, Limit, Torque, AMM")
    keyword_list = [k.strip().lower() for k in keywords.split(",")]
    
    st.divider()
    st.header("📜 Session History")
    if st.session_state.history:
        for item in reversed(st.session_state.history[-5:]):
            with st.expander(f"{item['time']} - {item['files'][:15]}..."):
                st.write(item['result'][:300] + "...")
    else:
        st.caption("No audits in this session yet.")

    if st.button("🗑️ Clear All Data"):
        st.session_state.batch_log = []
        st.session_state.history = []
        st.rerun()

if pwd_input != APP_PASSWORD:
    st.title("🔍 Smart-Diff Pro")
    st.warning("Please enter the App Password in the sidebar to unlock.")
    st.stop()

# --- 4. MAIN INTERFACE ---
st.title("🔍 Smart-Diff Pro")
st.caption("Fleet-Master 2026 | Safety Guard | Optimized for Ethiopian MRO")

# SECTION 1: MASTER (Verified Selection Logic)
st.subheader("1. Master Source (Source of Truth)")
m_t1, m_t2 = st.tabs(["📤 Upload Master", "📂 Local Master Folder"])
master_to_use = None

with m_t1:
    m_up = st.file_uploader("Upload Master Document", type=['pdf', 'docx'])
    if m_up: master_to_use = m_up

with m_t2:
    m_path = st.text_input("Master Folder Path (C:\\...):", key="m_path_input")
    if m_path:
        clean_m = os.path.normpath(m_path.strip().strip('"'))
        if os.path.exists(clean_m):
            m_files = [os.path.join(clean_m, f) for f in os.listdir(clean_m) if f.lower().endswith(('.pdf', '.docx'))]
            if m_files:
                m_sel = st.selectbox("Select Master File from Folder:", [os.path.basename(f) for f in m_files])
                master_to_use = next(f for f in m_files if os.path.basename(f) == m_sel)
                st.success(f"✅ Ready: {os.path.basename(master_to_use)}")
        else: st.error("Directory not found.")

# SECTION 2: REVISED (Synchronized Selection Logic)
st.subheader("2. Revised Documents (Batch Analysis)")
r_t1, r_t2 = st.tabs(["📤 Manual Batch Upload", "📂 Local Folder Watcher"])
rev_queue = []

with r_t1:
    r_up = st.file_uploader("Upload Revised Documents", type=['pdf', 'docx'], accept_multiple_files=True)
    if r_up: rev_queue.extend(r_up)

with r_t2:
    r_path = st.text_input("Revised Folder Path (C:\\...):", key="r_path_input")
    if r_path:
        clean_r = os.path.normpath(r_path.strip().strip('"'))
        if os.path.exists(clean_r):
            r_found = [os.path.join(clean_r, f) for f in os.listdir(clean_r) if f.lower().endswith(('.pdf', '.docx'))]
            if r_found:
                # NEW: Dropdown for Revised selection to mirror Master behavior
                r_choices = [os.path.basename(f) for f in r_found]
                selected_r = st.multiselect("Select Files to Include in Audit:", r_choices, default=r_choices)
                rev_queue.extend([f for f in r_found if os.path.basename(f) in selected_r])
                st.success(f"✅ {len(rev_queue)} files added to queue.")
        else: st.error("Directory not found.")

# SECTION 3: PREVIEWER (Verified Side-by-Side)
if master_to_use and rev_queue:
    if st.toggle("👁️ Show Side-by-Side Previewer"):
        v1, v2 = st.columns(2)
        v1.text_area("Master Content", extract_text(master_to_use), height=250)
        if isinstance(master_to_use, str):
            v1.button("📂 Open Master on PC", on_click=open_local_file, args=(master_to_use,))
            
        r_names = [f.name if hasattr(f, 'name') else os.path.basename(f) for f in rev_queue]
        sel_r = v2.selectbox("Select Revised to Preview:", r_names)
        target_r = rev_queue[r_names.index(sel_r)]
        v2.text_area(f"Revised Preview: {sel_r}", extract_text(target_r), height=250)
        if isinstance(target_r, str):
            v2.button(f"📂 Open '{sel_r}'", key="prev_op", on_click=open_local_file, args=(target_r,))

st.divider()

# --- 5. EXECUTION ENGINE (Verified Multi-Feature Process) ---
if st.button("🚀 Run Full Safety Audit"):
    if master_to_use and rev_queue:
        t_master = extract_text(master_to_use)[:35000]
        prog = st.progress(0)
        
        for i, file in enumerate(rev_queue):
            fname = file.name if hasattr(file, 'name') else os.path.basename(file)
            prog.progress(int(((i+1)/len(rev_queue))*100))
            
            with st.status(f"Auditing {fname}..."):
                t_rev = extract_text(file)[:35000]
                report = run_audit(t_master, t_rev, api_key, focus_mode)
                red_html = Redlines(t_master, t_rev).output_markdown
                score = fuzz.token_set_ratio(t_master, t_rev)
                
                # Verified Safety Guard Logic
                has_num_change = any(c.isdigit() for c in report)
                found_alerts = [k.upper() for k in keyword_list if k in report.lower()]
                risk_lvl = "🚨 CRITICAL" if (has_num_change and found_alerts) else "🚨 HIGH" if has_num_change else "🟢 LOW"
                
                st.session_state.batch_log.append({
                    "Timestamp": datetime.now().strftime("%H:%M"),
                    "File": fname, 
                    "Match %": f"{score}%", 
                    "Risk": risk_lvl, 
                    "Alerts": ", ".join(found_alerts)
                })
                st.session_state.history.append({"time": datetime.now().strftime("%H:%M"), "files": fname, "result": report})

                # Individual Result Layout
                st.subheader(f"📄 Audit Result: {fname}")
                c1, c2 = st.columns(2)
                if isinstance(file, str):
                    c1.button(f"📂 Open {fname}", key=f"o_btn_{i}", on_click=open_local_file, args=(file,))
                    if c2.button(f"📦 Archive {fname}", key=f"a_btn_{i}"):
                        if archive_file(file): st.success(f"Moved to Audited_Results")
                
                if found_alerts: 
                    st.markdown(f'<div class="keyword-alert">⚠️ Keywords Detected: {", ".join(found_alerts)}</div>', unsafe_allow_html=True)
                
                res_tab1, res_tab2 = st.tabs(["📊 AI Detailed Report", "🎨 Visual Redline (Track Changes)"])
                with res_tab1: st.markdown(report)
                with res_tab2: st.markdown(red_html, unsafe_allow_html=True)
                st.divider()
    else: 
        st.warning("Required: Select both a Master Document and at least one Revised Document.")

# --- 6. GLOBAL DASHBOARD & EXPORT ---
if st.session_state.batch_log:
    st.header("📊 Batch Audit Dashboard")
    df = pd.DataFrame(st.session_state.batch_log)
    
    m1, m2, m3 = st.columns(3)
    m1.metric("Files Processed", len(df))
    m2.metric("Critical Risks", len(df[df['Risk'].str.contains("🚨")]))
    m3.metric("Avg Match %", f"{int(df['Match %'].str.replace('%','').astype(int).mean())}%")
    
    st.dataframe(df, use_container_width=True)
    
    # Advanced Excel Export
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='AuditLog')
    st.download_button("📥 Download Final Excel Report", output.getvalue(), "MRO_Audit_Report.xlsx")